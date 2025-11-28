"""
Module xuất báo cáo Word
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os


def create_deduplication_report(
    result: dict,
    method_name: str,
    output_path: str,
    performance: dict = None
):
    """
    Tạo báo cáo Word từ kết quả phân cụm
    
    Args:
        result: Dict từ process_clustering()
        method_name: Tên phương pháp (SimHash, MinHash, FAISS)
        output_path: Đường dẫn lưu file
        performance: Dict hiệu năng (thời gian, số văn bản, v.v.)
    """
    
    doc = Document()
    
    # ===== TIÊU ĐỀ =====
    title = doc.add_heading('BÁO CÁO PHÁT HIỆN VÀ LOẠI BỎ TRÙNG LẶP', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_paragraph(f'Phương pháp: {method_name}')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.runs[0]
    run.font.size = Pt(12)
    run.font.bold = True
    
    time_para = doc.add_paragraph(f'Ngày: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    time_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()  # Khoảng cách
    
    # ===== THỐNG KÊ =====
    doc.add_heading('1. THỐNG KÊ TỔNG QUAN', 1)
    
    stats = result.get('stats', {})
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    stats_items = [
        ('Tổng số văn bản', stats.get('total_docs', 0)),
        ('Số cụm phát hiện', stats.get('n_clusters', 0)),
        ('Văn bản bị loại', stats.get('n_removed', 0)),
        ('Văn bản giữ lại', stats.get('n_kept', 0)),
        ('Tỷ lệ loại bỏ', f"{stats.get('removal_rate', 0):.1%}"),
        ('Số cặp tương tự', stats.get('n_pairs', 0)),
    ]
    
    for label, value in stats_items:
        row = table.add_row()
        cell_label = row.cells[0]
        cell_value = row.cells[1]
        
        cell_label.text = label
        cell_value.text = str(value)
        
        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    doc.add_paragraph()
    
    # ===== HIỆU NĂNG =====
    if performance:
        doc.add_heading('2. HIỆU NĂNG', 1)
        
        perf_table = doc.add_table(rows=1, cols=2)
        perf_table.style = 'Light Grid Accent 1'
        
        for key, value in performance.items():
            row = perf_table.add_row()
            row.cells[0].text = str(key)
            row.cells[1].text = str(value)
        
        doc.add_paragraph()
    
    # ===== CHI TIẾT CÁC CỤM =====
    doc.add_heading('3. CHI TIẾT CÁC CỤM TRÙNG LẶP', 1)
    
    clusters = result.get('clusters', {})
    
    if not clusters:
        doc.add_paragraph('Không tìm thấy cụm trùng lặp')
    else:
        for cluster_idx, (cluster_id, cluster_info) in enumerate(clusters.items(), 1):
            # Tiêu đề cụm
            doc.add_heading(f'Cụm #{cluster_idx}', 2)
            
            size = cluster_info.get('size', 0)
            doc.add_paragraph(f'Số lượng văn bản: {size}')
            
            # Danh sách documents
            documents = cluster_info.get('documents', [])
            
            # Văn bản đại diện
            representatives = [d for d in documents if d.get('is_representative')]
            if representatives:
                rep_para = doc.add_paragraph()
                rep_para.paragraph_format.left_indent = Inches(0.3)
                
                run = rep_para.add_run('✓ ĐẠI DIỆN (GIỮ LẠI)')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 128, 0)
                run.font.size = Pt(11)
                
                for rep_doc in representatives:
                    text = rep_doc.get('text', '')
                    doc_id = rep_doc.get('id', 'N/A')
                    
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.add_run(f'ID: {doc_id}\n').font.size = Pt(10)
                    p.add_run(text).font.size = Pt(10)
            
            # Văn bản trùng lặp
            duplicates = [d for d in documents if not d.get('is_representative')]
            if duplicates:
                dup_para = doc.add_paragraph()
                dup_para.paragraph_format.left_indent = Inches(0.3)
                
                run = dup_para.add_run('✗ TRÙNG LẶP (LOẠI BỎ)')
                run.font.bold = True
                run.font.color.rgb = RGBColor(200, 0, 0)
                run.font.size = Pt(11)
                
                for dup_doc in duplicates:
                    text = dup_doc.get('text', '')
                    doc_id = dup_doc.get('id', 'N/A')
                    
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    
                    id_run = p.add_run(f'ID: {doc_id}\n')
                    id_run.font.size = Pt(10)
                    id_run.font.color.rgb = RGBColor(100, 100, 100)
                    
                    text_run = p.add_run(text)
                    text_run.font.size = Pt(10)
                    text_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # Đường phân cách
            doc.add_paragraph('─' * 80)
    
    # ===== KẾT LUẬN =====
    doc.add_heading('4. KẾT LUẬN', 1)
    
    conclusion = f"""
Báo cáo này cho thấy kết quả phát hiện và loại bỏ các văn bản trùng lặp sử dụng phương pháp {method_name}.

Các văn bản được phân thành {stats.get('n_clusters', 0)} cụm, trong đó:
- {stats.get('n_kept', 0)} văn bản được giữ lại (đại diện)
- {stats.get('n_removed', 0)} văn bản bị loại bỏ (trùng lặp)

Tỷ lệ loại bỏ: {stats.get('removal_rate', 0):.1%}
    """
    
    doc.add_paragraph(conclusion.strip())
    
    # Lưu file
    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
    doc.save(output_path)
    
    print(f"✓ Báo cáo đã lưu: {output_path}")


if __name__ == '__main__':
    # Test
    test_result = {
        'stats': {
            'total_docs': 10,
            'n_clusters': 3,
            'n_removed': 5,
            'n_kept': 5,
            'n_pairs': 5,
            'removal_rate': 0.5
        },
        'clusters': {
            0: {
                'size': 2,
                'representative': 0,
                'documents': [
                    {'id': 0, 'text': 'Văn bản 1', 'is_representative': True},
                    {'id': 1, 'text': 'Văn bản 1 sửa', 'is_representative': False}
                ]
            }
        }
    }
    
    create_deduplication_report(
        test_result,
        'FAISS',
        'test_report.docx',
        {'Thời gian': '2.5s', 'Số văn bản': 10}
    )