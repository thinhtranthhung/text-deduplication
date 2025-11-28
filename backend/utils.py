"""
Module xử lý file input - hỗ trợ TXT, DOC, DOCX, CSV, JSON
"""
import os
import csv
import json
from io import BytesIO
from werkzeug.datastructures import FileStorage


def extract_text_from_file(file: FileStorage) -> list:
    """
    Đọc nội dung từ file và trả về danh sách văn bản
    
    Args:
        file: FileStorage object từ Flask
    
    Returns:
        List các văn bản đã được tách ra từ file
    
    Raises:
        Exception: Nếu định dạng file không hỗ trợ hoặc file trống
    """
    
    filename = file.filename.lower()
    texts = []
    
    try:
        # ========== TXT FILE ==========
        if filename.endswith('.txt'):
            content = file.read().decode('utf-8', errors='ignore')
            # Tách theo dòng, lọc dòng trống
            texts = [line.strip() for line in content.split('\n') if line.strip()]
        
        # ========== CSV FILE ==========
        elif filename.endswith('.csv'):
            content = file.read().decode('utf-8', errors='ignore')
            reader = csv.reader(content.splitlines())
            for row in reader:
                # Gộp tất cả cột hoặc lấy cột đầu tiên
                text = ' '.join([col.strip() for col in row if col.strip()])
                if text:
                    texts.append(text)
        
        # ========== JSON FILE ==========
        elif filename.endswith('.json'):
            content = file.read().decode('utf-8', errors='ignore')
            data = json.loads(content)
            
            if isinstance(data, list):
                for item in data:
                    if isinstance(item, dict):
                        # Tìm key chứa text
                        for key in ['content', 'text', 'body', 'message', 'data']:
                            if key in item and isinstance(item[key], str):
                                text = item[key].strip()
                                if text:
                                    texts.append(text)
                                break
                    elif isinstance(item, str):
                        text = item.strip()
                        if text:
                            texts.append(text)
            elif isinstance(data, dict):
                for value in data.values():
                    if isinstance(value, str):
                        text = value.strip()
                        if text:
                            texts.append(text)
        
        # ========== DOC/DOCX FILE ==========
        elif filename.endswith(('.doc', '.docx')):
            from docx import Document
            
            file.seek(0)
            doc_bytes = BytesIO(file.read())
            doc = Document(doc_bytes)
            
            # Lấy text từ paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    texts.append(text)
            
            # Nếu có tables, lấy text từ bảng
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            texts.append(text)
        
        else:
            raise ValueError(f"Định dạng file '{filename}' không được hỗ trợ")
        
        # Kiểm tra dữ liệu hợp lệ
        if not texts:
            raise ValueError("File không chứa nội dung hợp lệ")
        
        if len(texts) < 2:
            raise ValueError(f"File phải chứa ít nhất 2 văn bản (hiện có {len(texts)})")
        
        print(f"✓ Đã đọc {len(texts)} văn bản từ file")
        return texts
    
    except Exception as e:
        raise Exception(f"Lỗi khi xử lý file: {str(e)}")